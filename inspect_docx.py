from docx import Document
import os

# Find the docx file dynamically
files = [f for f in os.listdir('.') if f.endswith('.docx') and not f.startswith('~')]
if not files:
    print("No DOCX file found.")
    exit()

doc_path = os.path.abspath(files[0])
print(f"Analyzing file: {doc_path}")

try:
    doc = Document(doc_path)
    print(f"Document has {len(doc.paragraphs)} paragraphs.")
    print(f"Document has {len(doc.tables)} tables.")
    
    print("-" * 30)
    total_chars = 0
    total_words = 0
    for p in doc.paragraphs:
        total_chars += len(p.text)
        total_words += len(p.text.split())
    
    print(f"Total Character Count: {total_chars}")
    print(f"Total Word Count: {total_words}")
    print("-" * 30)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text and (len(text) < 100) and (text.isupper() or text[0].isdigit() or "INTRODUCTION" in text.upper() or "ABSTRACT" in text.upper() or "CONCLUSION" in text.upper() or "REFERENCE" in text.upper()):
             print(f"Index {i}: {text}")
except Exception as e:
    print(f"Error reading docx: {e}")

